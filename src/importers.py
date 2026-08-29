"""从 TXT / Excel / Word / Markdown 导入英语生词。"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Iterable


WORD_RE = re.compile(r"^[A-Za-z][A-Za-z\-']{0,48}$")
ENGLISH_TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z\-']{1,47}")
PHONETIC_RE = re.compile(r"\[[^\]]*\]|/[^/\s][^/]*/")
HEADER_WORDS = {
    "word", "words", "vocabulary", "vocab", "english", "en", "term", "terms",
    "meaning", "chinese", "translation", "note", "notes", "pos", "level",
    "spell", "spelling", "unit", "lesson", "page", "no", "id", "index",
}
POS_ABBR = {
    "n", "v", "vt", "vi", "adj", "adv", "prep", "conj", "pron", "art", "num",
    "pl", "sing", "sth", "sb", "etc", "cf",
}
HEADER_KEYWORDS = (
    "word", "单词", "词汇", "词语", "生词", "英文", "英语", "vocab",
    "vocabulary", "english", "spell", "拼写", "term",
)
SKIP_CELL_VALUES = {"", "nan", "none", "null", "nat", "n/a", "#n/a", "-", "—", "－"}


def _cell_text(raw: object) -> str:
    if raw is None:
        return ""
    if isinstance(raw, float) and raw != raw:  # NaN
        return ""
    text = str(raw).strip()
    if text.lower() in SKIP_CELL_VALUES:
        return ""
    # Excel 有时把数字读成 1.0
    if re.fullmatch(r"\d+\.0", text):
        text = text[:-2]
    return text


def _normalize_word(raw: object) -> str | None:
    """从单元格提取一个英文生词（支持「abandon 放弃」「放弃 abandon」「1. apple」等）。"""
    text = _cell_text(raw)
    if not text:
        return None

    text = re.sub(r"^[\d]+[\.\)、．]\s*", "", text)
    text = re.sub(r"^[-*•]\s*", "", text)
    text = PHONETIC_RE.sub(" ", text)
    text = text.replace("\u3000", " ")

    # 优先：以英文开头
    m = re.match(r"^([A-Za-z][A-Za-z\-']*)", text)
    if m:
        word = m.group(1).lower()
        # 单字母（除 a/i 外无意义）直接跳过
        if (
            len(word) >= 2
            and WORD_RE.match(word)
            and word not in HEADER_WORDS
            and word not in POS_ABBR
        ):
            return word

    # 否则：从整格中找最像单词的英文 token（跳过词性缩写）
    candidates: list[str] = []
    for token in ENGLISH_TOKEN_RE.findall(text):
        w = token.lower().strip("-'")
        if len(w) < 2:
            continue
        if w in HEADER_WORDS or w in POS_ABBR:
            continue
        if not WORD_RE.match(w):
            continue
        candidates.append(w)

    if not candidates:
        return None
    # 取最长的，避免抓到 vt/to 等残留；同长度取第一个
    return max(candidates, key=lambda w: (len(w), -candidates.index(w)))


def _dedupe(words: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for w in words:
        if w and w not in seen:
            seen.add(w)
            result.append(w)
    return result


def import_txt(path: Path | str, content: bytes | str | None = None) -> list[str]:
    if content is None:
        text = Path(path).read_text(encoding="utf-8", errors="ignore")
    elif isinstance(content, bytes):
        text = content.decode("utf-8", errors="ignore")
    else:
        text = content
    words: list[str] = []
    for line in text.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        parts = re.split(r"[,;\t|]+", line) if ("," in line or "\t" in line or ";" in line) else [line]
        for part in parts:
            w = _normalize_word(part)
            if w:
                words.append(w)
    return _dedupe(words)


def import_markdown(path: Path | str, content: bytes | str | None = None) -> list[str]:
    if content is None:
        text = Path(path).read_text(encoding="utf-8", errors="ignore")
    elif isinstance(content, bytes):
        text = content.decode("utf-8", errors="ignore")
    else:
        text = content
    text = re.sub(r"```.*?```", "\n", text, flags=re.S)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.M)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"[*`_>~]", "", text)
    words: list[str] = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        line = re.sub(r"^[-*+]\s+", "", line)
        line = re.sub(r"^\d+[\.\)]\s+", "", line)
        if "|" in line:
            cells = [
                c.strip()
                for c in line.split("|")
                if c.strip() and not re.fullmatch(r":?-+:?", c.strip())
            ]
            for cell in cells:
                w = _normalize_word(cell)
                if w:
                    words.append(w)
            continue
        w = _normalize_word(line)
        if w:
            words.append(w)
    return _dedupe(words)


def _read_excel_sheets(content: bytes | None, path: Path | str | None) -> dict[str, list[list[object]]]:
    """读取全部工作表为 {表名: 行列表}，兼容 xlsx / xlsm（openpyxl）与 xls（xlrd）。"""
    last_error: Exception | None = None

    # 1) xlsx / xlsm：openpyxl
    try:
        from openpyxl import load_workbook

        source = BytesIO(content) if content is not None else str(path)
        wb = load_workbook(source, read_only=True, data_only=True)
        sheets = {
            ws.title: [list(row) for row in ws.iter_rows(values_only=True)]
            for ws in wb.worksheets
        }
        wb.close()
        if sheets:
            return sheets
    except Exception as e:  # noqa: BLE001 — 失败时回退 xls 引擎
        last_error = e

    # 2) 旧版 xls：xlrd
    try:
        import xlrd

        if content is not None:
            book = xlrd.open_workbook(file_contents=content)
        else:
            book = xlrd.open_workbook(str(path))
        sheets = {
            sheet.name: [sheet.row_values(r) for r in range(sheet.nrows)]
            for sheet in book.sheets()
        }
        if sheets:
            return sheets
    except Exception as e:  # noqa: BLE001
        last_error = e

    hint = "请确认文件为 .xlsx / .xls，或另存为 Excel 工作簿后再试。"
    raise ValueError(f"无法读取 Excel 文件。{hint} 详情：{last_error}")


def _header_match_score(cell: object) -> int:
    text = _cell_text(cell).lower()
    if not text:
        return 0
    score = 0
    for key in HEADER_KEYWORDS:
        if key in text:
            score += 3 if key in {"单词", "生词", "英文", "word", "vocabulary"} else 2
    return score


def _words_from_column(rows: list[list[object]], col: int, start_row: int = 0) -> list[str]:
    words: list[str] = []
    for row in rows[start_row:]:
        if col < len(row):
            w = _normalize_word(row[col])
            if w:
                words.append(w)
    return words


def _pick_words_from_rows(rows: list[list[object]]) -> list[str]:
    rows = [r for r in rows if any(_cell_text(c) for c in r)]
    if not rows:
        return []

    width = max(len(r) for r in rows)
    rows = [list(r) + [""] * (width - len(r)) for r in rows]

    # 1) 表头关键字命中的列优先
    header_scores = sorted(
        ((_header_match_score(c), i) for i, c in enumerate(rows[0])),
        reverse=True,
    )
    if header_scores and header_scores[0][0] > 0:
        best_score = header_scores[0][0]
        collected: list[str] = []
        for score, col in header_scores:
            if score < best_score:
                break
            collected.extend(_words_from_column(rows, col, start_row=1))
        if collected:
            return _dedupe(collected)

    # 2) 无明确表头：选「英词密度」最高的列
    best_words: list[str] = []
    best_score = -1
    for col in range(width):
        col_words = _words_from_column(rows, col, start_row=0)
        if len(col_words) > best_score:
            best_score = len(col_words)
            best_words = col_words
    if best_words:
        return _dedupe(best_words)

    # 3) 兜底：整表逐格扫描
    words: list[str] = []
    for row in rows:
        for val in row:
            w = _normalize_word(val)
            if w:
                words.append(w)
    return _dedupe(words)


def import_excel(path: Path | str, content: bytes | None = None) -> list[str]:
    sheets = _read_excel_sheets(content, path)
    all_words: list[str] = []
    for _name, rows in sheets.items():
        all_words.extend(_pick_words_from_rows(rows))
    return _dedupe(all_words)


def import_word(path: Path | str, content: bytes | None = None) -> list[str]:
    from docx import Document

    if content is not None:
        doc = Document(BytesIO(content))
    else:
        doc = Document(str(path))
    words: list[str] = []
    for para in doc.paragraphs:
        for line in para.text.splitlines():
            w = _normalize_word(line)
            if w:
                words.append(w)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                w = _normalize_word(cell.text)
                if w:
                    words.append(w)
    return _dedupe(words)


def import_vocab_file(
    filename: str,
    content: bytes | None = None,
    path: Path | str | None = None,
) -> list[str]:
    """根据扩展名自动选择导入器。"""
    name = (filename or str(path or "")).lower()
    suffix = Path(name).suffix

    if suffix in {".txt", ".csv"}:
        return import_txt(path or name, content)
    if suffix in {".md", ".markdown"}:
        return import_markdown(path or name, content)
    if suffix in {".xlsx", ".xls", ".xlsm"}:
        return import_excel(path or name, content)
    if suffix in {".docx"}:
        return import_word(path or name, content)
    if suffix == ".doc":
        raise ValueError("暂不支持旧版 .doc，请另存为 .docx 后再导入。")
    return import_txt(path or name, content)


def parse_manual_input(text: str) -> list[str]:
    """解析用户手动粘贴的单词列表。"""
    return import_txt("manual.txt", text)
